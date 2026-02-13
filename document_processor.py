"""
document_processor.py
Core find-and-replace logic for Word documents (.docx).

Handles:
- Cross-run text matching (finds text that spans multiple runs)
- Positional replacement (replace individual occurrences selectively)
- Tables, headers, footers, and body paragraphs
- Formatting preservation during replacement
"""

import os
import shutil
from dataclasses import dataclass, field
from typing import List, Optional, Tuple
from docx import Document
from docx.opc.exceptions import PackageNotFoundError


@dataclass
class Match:
    """Represents a single find-text occurrence in a document."""
    file_path: str
    location_type: str          # 'body', 'table', 'header', 'footer'
    location_detail: str        # e.g., 'Paragraph 5', 'Table 2, Row 1, Cell 3, Para 1'
    paragraph_id: int           # Unique ID for this paragraph across the document
    match_index: int            # Which occurrence within this paragraph (0-based)
    char_offset: int            # Character offset within the paragraph's full text
    context_before: str         # ~50 chars before match
    match_text: str             # The matched text
    context_after: str          # ~50 chars after match
    is_selected: bool = True    # User's checkbox state

    @property
    def display_context(self) -> str:
        """Format the context for display in the GUI."""
        before = self.context_before
        after = self.context_after
        if len(before) > 50:
            before = "..." + before[-47:]
        if len(after) > 50:
            after = after[:47] + "..."
        return f"{before}[{self.match_text}]{after}"


@dataclass
class FileResult:
    """All matches for a single file."""
    file_path: str
    file_name: str
    matches: List[Match] = field(default_factory=list)
    error: Optional[str] = None

    @property
    def match_count(self) -> int:
        return len(self.matches)

    @property
    def selected_count(self) -> int:
        return sum(1 for m in self.matches if m.is_selected)


def _collect_paragraphs(doc: Document, file_path: str) -> List[Tuple[object, str, str, int]]:
    """
    Collect all paragraphs from body, tables, headers, and footers.
    
    Returns list of tuples: (paragraph_obj, location_type, location_detail, paragraph_id)
    """
    paragraphs = []
    pid = 0

    # Body paragraphs
    for i, para in enumerate(doc.paragraphs):
        paragraphs.append((para, 'body', f'Paragraph {i + 1}', pid))
        pid += 1

    # Table paragraphs
    for t_idx, table in enumerate(doc.tables):
        for r_idx, row in enumerate(table.rows):
            for c_idx, cell in enumerate(row.cells):
                for p_idx, para in enumerate(cell.paragraphs):
                    detail = f'Table {t_idx + 1}, Row {r_idx + 1}, Cell {c_idx + 1}'
                    if len(cell.paragraphs) > 1:
                        detail += f', Para {p_idx + 1}'
                    paragraphs.append((para, 'table', detail, pid))
                    pid += 1

    # Headers and footers
    for s_idx, section in enumerate(doc.sections):
        section_label = f'Section {s_idx + 1} ' if len(doc.sections) > 1 else ''

        for header_attr, label in [('header', 'Header'), ('first_page_header', 'First Page Header'),
                                    ('even_page_header', 'Even Page Header')]:
            try:
                header = getattr(section, header_attr)
                if header and header.paragraphs:
                    for p_idx, para in enumerate(header.paragraphs):
                        detail = f'{section_label}{label}'
                        if len(header.paragraphs) > 1:
                            detail += f', Para {p_idx + 1}'
                        paragraphs.append((para, 'header', detail, pid))
                        pid += 1
            except Exception:
                pass

        for footer_attr, label in [('footer', 'Footer'), ('first_page_footer', 'First Page Footer'),
                                    ('even_page_footer', 'Even Page Footer')]:
            try:
                footer = getattr(section, footer_attr)
                if footer and footer.paragraphs:
                    for p_idx, para in enumerate(footer.paragraphs):
                        detail = f'{section_label}{label}'
                        if len(footer.paragraphs) > 1:
                            detail += f', Para {p_idx + 1}'
                        paragraphs.append((para, 'footer', detail, pid))
                        pid += 1
            except Exception:
                pass

    return paragraphs


def _get_context(full_text: str, start: int, length: int, context_chars: int = 50) -> Tuple[str, str]:
    """Extract context before and after a match position."""
    before_start = max(0, start - context_chars)
    after_end = min(len(full_text), start + length + context_chars)

    context_before = full_text[before_start:start]
    context_after = full_text[start + length:after_end]

    if before_start > 0:
        context_before = "..." + context_before
    if after_end < len(full_text):
        context_after = context_after + "..."

    return context_before, context_after


def _find_all_occurrences(text: str, find_text: str, case_sensitive: bool = True) -> List[int]:
    """Find all character offsets of find_text in text."""
    offsets = []
    search_text = text if case_sensitive else text.lower()
    search_term = find_text if case_sensitive else find_text.lower()

    start = 0
    while True:
        idx = search_text.find(search_term, start)
        if idx == -1:
            break
        offsets.append(idx)
        start = idx + 1
    return offsets


def scan_documents(folder_path: str, find_text: str, case_sensitive: bool = True) -> List[FileResult]:
    """
    Scan all .docx files in a folder for occurrences of find_text.
    
    Returns a list of FileResult objects (one per file), each containing
    Match objects for every occurrence found.
    """
    if not find_text:
        return []

    results = []

    # Get all .docx files (non-recursive), skip temp files
    try:
        docx_files = sorted([
            os.path.join(folder_path, f)
            for f in os.listdir(folder_path)
            if f.lower().endswith('.docx') and not f.startswith('~$')
        ])
    except PermissionError as e:
        return [FileResult(file_path=folder_path, file_name=os.path.basename(folder_path),
                           error=f"Permission denied: {e}")]

    if not docx_files:
        return []

    for file_path in docx_files:
        file_result = FileResult(
            file_path=file_path,
            file_name=os.path.basename(file_path)
        )

        try:
            doc = Document(file_path)
        except PackageNotFoundError:
            file_result.error = "File appears to be corrupted or is not a valid .docx file"
            results.append(file_result)
            continue
        except PermissionError:
            file_result.error = "Permission denied (file may be open in another application)"
            results.append(file_result)
            continue
        except Exception as e:
            file_result.error = f"Error opening file: {str(e)}"
            results.append(file_result)
            continue

        # Collect all paragraphs from all locations
        all_paragraphs = _collect_paragraphs(doc, file_path)

        for para_obj, loc_type, loc_detail, para_id in all_paragraphs:
            full_text = para_obj.text
            if not full_text:
                continue

            offsets = _find_all_occurrences(full_text, find_text, case_sensitive)

            for match_idx, char_offset in enumerate(offsets):
                context_before, context_after = _get_context(
                    full_text, char_offset, len(find_text)
                )

                match = Match(
                    file_path=file_path,
                    location_type=loc_type,
                    location_detail=loc_detail,
                    paragraph_id=para_id,
                    match_index=match_idx,
                    char_offset=char_offset,
                    context_before=context_before,
                    match_text=find_text if case_sensitive else full_text[char_offset:char_offset + len(find_text)],
                    context_after=context_after,
                )
                file_result.matches.append(match)

        if file_result.matches or file_result.error:
            results.append(file_result)

    return results


def _replace_in_paragraph(paragraph, find_text: str, replace_text: str,
                          target_offsets: List[int], case_sensitive: bool = True) -> int:
    """
    Replace specific occurrences of find_text in a paragraph, handling cross-run matches.
    
    This is the core replacement algorithm:
    1. Build a map of (char_offset -> run_index, offset_within_run) for every character
    2. For each target occurrence, find which runs contain it
    3. Modify those runs to perform the replacement while preserving formatting
    
    Returns the number of replacements made.
    """
    runs = paragraph.runs
    if not runs:
        return 0

    # Build character-to-run mapping
    # char_map[i] = (run_index, offset_within_run)
    char_map = []
    for run_idx, run in enumerate(runs):
        for char_idx in range(len(run.text)):
            char_map.append((run_idx, char_idx))

    full_text = paragraph.text
    find_len = len(find_text)
    replace_count = 0

    # Process replacements in reverse order to preserve offsets
    for offset in sorted(target_offsets, reverse=True):
        # Validate the match still exists at this offset
        if offset + find_len > len(full_text):
            continue

        actual_text = full_text[offset:offset + find_len]
        if case_sensitive and actual_text != find_text:
            continue
        if not case_sensitive and actual_text.lower() != find_text.lower():
            continue

        if offset >= len(char_map) or offset + find_len - 1 >= len(char_map):
            continue

        # Find which runs are involved
        start_run_idx, start_char_idx = char_map[offset]
        end_run_idx, end_char_idx = char_map[offset + find_len - 1]

        if start_run_idx == end_run_idx:
            # Simple case: match is within a single run
            run = runs[start_run_idx]
            before = run.text[:start_char_idx]
            after = run.text[start_char_idx + find_len:]
            run.text = before + replace_text + after
        else:
            # Cross-run match: text spans multiple runs
            # Strategy: put replacement text in the first run, clear matched text from others

            # First run: keep text before match, add replacement
            first_run = runs[start_run_idx]
            before_text = first_run.text[:start_char_idx]
            first_run.text = before_text + replace_text

            # Middle runs (if any): clear entirely
            for mid_run_idx in range(start_run_idx + 1, end_run_idx):
                runs[mid_run_idx].text = ""

            # Last run: remove matched portion, keep text after
            last_run = runs[end_run_idx]
            last_run.text = last_run.text[end_char_idx + 1:]

        replace_count += 1

        # Rebuild char_map and full_text for next iteration (since we modified runs)
        char_map = []
        for run_idx, run in enumerate(runs):
            for char_idx in range(len(run.text)):
                char_map.append((run_idx, char_idx))
        full_text = "".join(run.text for run in runs)

    return replace_count


def apply_changes(matches: List[Match], replace_text: str,
                  create_backups: bool = True, case_sensitive: bool = True,
                  progress_callback=None) -> dict:
    """
    Apply selected replacements to documents.
    
    Args:
        matches: List of Match objects (only those with is_selected=True will be applied)
        replace_text: The replacement text
        create_backups: Whether to create .docx.bak backup files
        case_sensitive: Whether matching is case-sensitive
        progress_callback: Optional callable(file_name, file_index, total_files)
    
    Returns:
        dict with keys: 'total_replaced', 'files_modified', 'errors'
    """
    selected = [m for m in matches if m.is_selected]
    if not selected:
        return {'total_replaced': 0, 'files_modified': 0, 'errors': []}

    # Group by file
    files_dict = {}
    for match in selected:
        if match.file_path not in files_dict:
            files_dict[match.file_path] = []
        files_dict[match.file_path].append(match)

    total_replaced = 0
    files_modified = 0
    errors = []

    total_files = len(files_dict)

    for file_idx, (file_path, file_matches) in enumerate(files_dict.items()):
        file_name = os.path.basename(file_path)

        if progress_callback:
            progress_callback(file_name, file_idx, total_files)

        # Create backup
        if create_backups:
            backup_path = file_path + '.bak'
            try:
                shutil.copy2(file_path, backup_path)
            except Exception as e:
                errors.append(f"Could not create backup for {file_name}: {e}")
                continue

        try:
            doc = Document(file_path)
        except Exception as e:
            errors.append(f"Could not open {file_name}: {e}")
            continue

        # Re-collect paragraphs (same order as during scanning)
        all_paragraphs = _collect_paragraphs(doc, file_path)

        # Group matches by paragraph_id
        para_matches = {}
        for match in file_matches:
            if match.paragraph_id not in para_matches:
                para_matches[match.paragraph_id] = []
            para_matches[match.paragraph_id].append(match)

        file_replaced = 0

        for para_obj, loc_type, loc_detail, para_id in all_paragraphs:
            if para_id not in para_matches:
                continue

            target_offsets = [m.char_offset for m in para_matches[para_id]]
            count = _replace_in_paragraph(
                para_obj, 
                para_matches[para_id][0].match_text,
                replace_text, 
                target_offsets,
                case_sensitive
            )
            file_replaced += count

        if file_replaced > 0:
            try:
                doc.save(file_path)
                files_modified += 1
                total_replaced += file_replaced
            except PermissionError:
                errors.append(f"Permission denied saving {file_name} (file may be open)")
            except Exception as e:
                errors.append(f"Error saving {file_name}: {e}")

    if progress_callback:
        progress_callback("Done", total_files, total_files)

    return {
        'total_replaced': total_replaced,
        'files_modified': files_modified,
        'errors': errors
    }
